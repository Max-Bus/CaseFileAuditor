import os
import re
import docx
import pandas as pd
import datetime
import re
import os
import win32com.client as win32
from win32com.client import constants
from pandas import Timestamp
from pandas import NaT
from wordscrape import WordDocument
import zipfile
from lxml import etree
# Create list of paths to .doc files
rootdir = 'C:\\Users\\mkbcu\\OneDrive\\Desktop\\cases'
months =["january","february","march","april","may","june","july","august","september","october","november","december","jan","feb","mar","apr","jun","aug","sep","sept","oct","nov","dec"]
i =0
names = os.listdir('C:\\Users\\mkbcu\\OneDrive\\Desktop\\cases - Copy')
print(names)
def save_as_docx(path):
    # Opening MS Word
    word = win32.gencache.EnsureDispatch('Word.Application')
    print(path)
    doc = word.Documents.Open(path)
    doc.Activate()

    # Rename path with .docx
    new_file_abs = os.path.abspath(path)
    new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)

    # Save and Close
    word.ActiveDocument.SaveAs(
        new_file_abs, FileFormat=constants.wdFormatXMLDocument
    )
    doc.Close(False)


def regex(text):
    if isinstance(text,datetime.datetime) and not isinstance(text,Timestamp) and text is not NaT and text is not None:
        try:
            text = text.strftime('%m/%d/%Y')
        except:
            print(text)
            return text
    text = str(text)
    if(person != ""):
        for n in person:
            print(n)
            text = re.sub(n.strip(), "[REDACTED]", text,flags=re.I)
    for month in months:
        text = re.sub(r'' + month.strip() + ' \d{0,2}', "[REDACTED]", text, flags=re.I)
    # ssn's
    text = re.sub(r"\d{3}-?\d{2}-?\d{4}", "[REDACTED]", text,flags=re.I)
    # dates
    text = re.sub(r"\d{1,2}[-/]\d{1,2}[-/]\d{2,4}", "[REDACTED]", text,flags=re.I)
    # phone numbers
    text = re.sub(r"(\+\d{1,2}\s)?\(?\d{3}\)?[\s.-]\d{3}[\s.-]\d{4}", "[REDACTED]", text,flags=re.I)
    # email
    text = re.sub(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", "[REDACTED]", text,flags=re.I)
    return text

for subdir, dirs, files in os.walk(rootdir):
    person =""
    for file in files:
        i+=1
        firstname = None
        lastname = None
        for name in names:
            if name in subdir:
                person = name.split(",")
        print(os.path.join(subdir, file))

        '''
        if re.search(r'\.xlsx$',file) or re.search(r'\.xls$',file): #excel file
            # Give the location of the file
            df = pd.read_excel(os.path.join(subdir, file))
            for col in df.columns:
                for row in df.index.values:
                    #df[col][row] = regex(df[col][row])
            df.to_csv(os.path.join(subdir, re.sub(r'\.xlsx?$','.csv',file)),index=False)

        if re.search(r'\.doc$',file):
            save_as_docx(os.path.join(subdir, file))

        
        if re.search(r'\.txt$',file) or re.search(r'\.rtf$',file):
            text = ""
            with open(os.path.join(subdir, file), 'r') as readdoc:
                text = regex(readdoc.read())

            with open(os.path.join(subdir, file), 'w') as writedoc:
                writedoc.write(text)
        #'''
        if re.search(r'\.docx$',file):
            document = docx.Document(os.path.join(subdir, file))  # creating word reader object.
            for d in document.tables:
                for i in range(len(d.rows)):
                    for j in range(len(d.columns)):
                        try:
                            d.cell(i,j).text = regex(d.cell(i,j).text)
                        except:
                            print("cell failed to scrub")
            #for pgh in document.paragraphs:
                #print()
                #pgh.text = regex(pgh.text)

            document.save(os.path.join(subdir, file))

print(i)


